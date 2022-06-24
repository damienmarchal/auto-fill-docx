import sys
import os
import csv                                        # Needed to process python 
from importlib.machinery import SourceFileLoader  # Needed to import python module from filename
from subprocess import  Popen                     # Needed to call subproceess (libreoffice, pdftk)
from collections import OrderedDict               # We are processing such a structure
from docx import Document                         # Load, Modify and Save docx files. 


def convert_to_pdf(input_docx, out_folder):
    """Convert the docx into pdf files using openoffice"""
    p = Popen(["libreoffice", '--headless', '--convert-to', 'pdf', '--outdir',
               out_folder, input_docx])
    print(["libreoffice", '--convert-to', 'pdf', input_docx])
    p.communicate()

def concat_all_pdf_in_one(source_dir, out_file):
    p = ["pdftk"] + source_dir + ["cat", "output", out_file]
    print(p)
    p = Popen(p)
    p.communicate()

def substitute(template_file_path, variables, output_file_path):
    """Substitute the given variables in the template_file_path and save the resulting docx in output_file_path"""
    template_document = Document(template_file_path)

    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        for table in template_document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, variable_key, variable_value)

    template_document.save(output_file_path)


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)

if __name__ == '__main__':
    if len(sys.argv) != 4:
        print("USAGE: gen.py template.docx csv-database.csv data-to-token.py")

    template_file_path = sys.argv[1]
    database_file_path = sys.argv[2]
    datacvt_file_path = sys.argv[3]
    
    input_data = open(database_file_path)
    attendees = csv.DictReader(input_data, delimiter=',')
    print(attendees)
        
    # imports the module and evaluate it so we have a data producing system. 
    data_to_token = SourceFileLoader(os.path.basename(sys.argv[3]),sys.argv[3]).load_module().patterns

    
    
    basename = os.path.splitext(os.path.basename(template_file_path))[0]

    attendees = list(attendees)
    attendees = sorted(attendees, key=lambda x: x['NOM'])
    out_list = []
    for attendee in attendees:           
        rule = {}
        for token, get_data in data_to_token.items():
            rule[token] = get_data(attendee)
        print(rule) 
        
        out_file = basename+"-"+data_to_token["${ID}"](attendee)
        docx_file = "tmp/"+out_file+".docx"
        substitute(template_file_path, rule, docx_file)        
        convert_to_pdf(docx_file, "result") 
        out_list.append("result/"+out_file+".pdf")
    
    concat_all_pdf_in_one(out_list, basename+"-all-in-one.pdf")
        
